# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx"]
# ///

"""
从 Word 简历中提取文本内容，保留结构和格式信息。

用法：
    uv run scripts/extract_docx.py resume.docx
    uv run scripts/extract_docx.py resume.docx --output extracted.txt
    uv run scripts/extract_docx.py resume.docx --format structured
"""

import argparse
import json
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _tag(ns: str, local: str) -> str:
    return f"{{{ns}}}{local}"


def extract_via_python_docx(docx_path: Path) -> list[dict]:
    """用 python-docx 提取，保留段落样式和 run 格式。"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("错误：python-docx 不可用。请运行：uv sync", file=sys.stderr)
        sys.exit(1)

    doc = Document(docx_path)
    elements = []

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        style_name = para.style.name if para.style else ""
        level = None
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1

        is_bold = all(run.bold for run in para.runs if run.text.strip()) if para.runs else False

        is_list = False
        if para._element is not None:
            pPr = para._element.find(_tag(W_NS, "pPr"))
            if pPr is not None:
                numPr = pPr.find(_tag(W_NS, "numPr"))
                if numPr is not None:
                    is_list = True

        alignment = None
        if para.alignment is not None:
            try:
                alignment = para.alignment.name if hasattr(para.alignment, 'name') else str(para.alignment)
            except Exception:
                pass

        el = {
            "type": "heading" if level else ("list_item" if is_list else "paragraph"),
            "text": para.text.strip(),
            "style": style_name,
        }
        if level:
            el["level"] = level
        if is_bold and not level:
            el["bold"] = True
        if alignment and alignment.upper() == "CENTER":
            el["centered"] = True

        elements.append(el)

    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            if any(row_cells):
                rows_data.append(row_cells)
        if rows_data:
            elements.append({"type": "table", "rows": rows_data})

    return elements


def extract_via_xml(docx_path: Path) -> list[dict]:
    """直接解析 document.xml，不依赖 python-docx，能处理更多边界情况。"""
    elements = []

    with zipfile.ZipFile(docx_path, "r") as zf:
        if "word/document.xml" not in zf.namelist():
            print("错误：DOCX 中未找到 word/document.xml", file=sys.stderr)
            return elements

        with zf.open("word/document.xml") as f:
            tree = ET.parse(f)

    root = tree.getroot()

    heading_styles = {
        "Heading1": 1, "Heading2": 2, "Heading3": 3, "Heading4": 4,
        "heading 1": 1, "heading 2": 2, "heading 3": 3, "heading 4": 4,
        "标题 1": 1, "标题 2": 2, "标题 3": 3, "标题 4": 4,
        "Title": 0,
    }

    p_tag = _tag(W_NS, "p")
    pPr_tag = _tag(W_NS, "pPr")
    pStyle_tag = _tag(W_NS, "pStyle")
    numPr_tag = _tag(W_NS, "numPr")
    r_tag = _tag(W_NS, "r")
    rPr_tag = _tag(W_NS, "rPr")
    t_tag = _tag(W_NS, "t")
    b_tag = _tag(W_NS, "b")
    jc_tag = _tag(W_NS, "jc")
    val_attr = _tag(W_NS, "val")
    tbl_tag = _tag(W_NS, "tbl")
    tr_tag = _tag(W_NS, "tr")
    tc_tag = _tag(W_NS, "tc")

    body = root.find(f".//{_tag(W_NS, 'body')}")
    if body is None:
        return elements

    for child in body:
        if child.tag == p_tag:
            text_parts = []
            for t_elem in child.findall(f".//{t_tag}"):
                if t_elem.text:
                    text_parts.append(t_elem.text)
            text = "".join(text_parts).strip()
            if not text:
                continue

            pPr = child.find(pPr_tag)
            style_name = ""
            level = None
            is_list = False
            is_bold = False
            is_centered = False

            if pPr is not None:
                pStyle = pPr.find(pStyle_tag)
                if pStyle is not None:
                    style_name = pStyle.get(val_attr, "")
                    level = heading_styles.get(style_name)

                if pPr.find(numPr_tag) is not None:
                    is_list = True

                jc = pPr.find(jc_tag)
                if jc is not None and jc.get(val_attr, "") == "center":
                    is_centered = True

            runs = child.findall(r_tag)
            bold_runs = 0
            total_runs = 0
            for run in runs:
                run_text = "".join(t.text or "" for t in run.findall(t_tag))
                if not run_text.strip():
                    continue
                total_runs += 1
                rPr = run.find(rPr_tag)
                if rPr is not None and rPr.find(b_tag) is not None:
                    bold_runs += 1

            if total_runs > 0 and bold_runs == total_runs:
                is_bold = True

            if level is None and is_bold and not is_list and len(text) < 30:
                level = 2

            el = {
                "type": "heading" if level is not None else ("list_item" if is_list else "paragraph"),
                "text": text,
            }
            if style_name:
                el["style"] = style_name
            if level is not None:
                el["level"] = level
            if is_bold and level is None:
                el["bold"] = True
            if is_centered:
                el["centered"] = True

            elements.append(el)

        elif child.tag == tbl_tag:
            rows_data = []
            for tr in child.findall(f".//{tr_tag}"):
                cells = []
                for tc in tr.findall(tc_tag):
                    cell_texts = []
                    for t_elem in tc.findall(f".//{t_tag}"):
                        if t_elem.text:
                            cell_texts.append(t_elem.text)
                    cells.append("".join(cell_texts).strip())
                if any(cells):
                    rows_data.append(cells)
            if rows_data:
                elements.append({"type": "table", "rows": rows_data})

    return elements


def elements_to_text(elements: list[dict]) -> str:
    """将结构化元素转为可读文本，用 Markdown 风格标记结构。"""
    lines = []
    for el in elements:
        etype = el["type"]
        if etype == "heading":
            level = el.get("level", 1)
            prefix = "#" * max(1, min(level, 4))
            lines.append(f"\n{prefix} {el['text']}")
        elif etype == "list_item":
            lines.append(f"  - {el['text']}")
        elif etype == "table":
            for row in el["rows"]:
                lines.append(" | ".join(row))
            lines.append("")
        else:
            text = el["text"]
            if el.get("bold"):
                text = f"**{text}**"
            if el.get("centered"):
                lines.append(f"[居中] {text}")
            else:
                lines.append(text)
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="从 Word 简历中提取文本")
    parser.add_argument("docx_path", type=Path, help="DOCX 文件路径")
    parser.add_argument("-o", "--output", type=Path, help="输出文件路径（默认输出到终端）")
    parser.add_argument(
        "--format",
        choices=["text", "structured", "json"],
        default="text",
        help="输出格式：text（带 Markdown 标记的可读文本）、structured/json（JSON 结构化数据）",
    )
    parser.add_argument(
        "--engine",
        choices=["auto", "python-docx", "xml"],
        default="auto",
        help="提取引擎：auto（先试 python-docx，失败用 xml）、python-docx、xml",
    )

    args = parser.parse_args()

    if not args.docx_path.exists():
        print(f"错误：文件不存在：{args.docx_path}", file=sys.stderr)
        sys.exit(1)

    if not args.docx_path.suffix.lower() == ".docx":
        print(f"错误：不是 DOCX 文件：{args.docx_path}", file=sys.stderr)
        sys.exit(1)

    if args.engine == "xml":
        elements = extract_via_xml(args.docx_path)
    elif args.engine == "python-docx":
        elements = extract_via_python_docx(args.docx_path)
    else:
        try:
            elements = extract_via_python_docx(args.docx_path)
        except Exception:
            print("python-docx 提取失败，回退到 XML 直接解析", file=sys.stderr)
            elements = extract_via_xml(args.docx_path)

    if args.format in ("structured", "json"):
        output = json.dumps(elements, indent=2, ensure_ascii=False)
    else:
        output = elements_to_text(elements)

    if args.output:
        args.output.write_text(output, encoding="utf-8")
        print(f"已写入：{args.output}（共 {len(elements)} 个元素）", file=sys.stderr)
    else:
        print(output)


if __name__ == "__main__":
    main()
