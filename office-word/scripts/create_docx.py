# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx"]
# ///

"""
从零生成标准排版的专利 .docx 文件。

用法：
    uv run scripts/create_docx.py content.json output.docx

仅当用户没有提供任何模板/原文件时使用。
如果用户有自己的模板，应使用 docx_edit.py replace 模式。

content.json 格式：
    {
      "claims": [
        {"number": 1, "text": "一种……方法，其特征在于……", "dependent": false},
        {"number": 2, "text": "根据权利要求1所述的方法，……", "dependent": true}
      ],
      "specification": {
        "field": "本发明涉及XX技术领域。",
        "background": ["段落1", "段落2"],
        "summary": ["段落1"],
        "drawings": ["图1为……"],
        "detailed": ["段落1", "段落2"]
      },
      "abstract": "摘要内容"
    }

可以只提供部分字段，比如只给 claims 就只生成权利要求书。

内置排版规范：
    纸张: A4 (21cm × 29.7cm)    页边距: 上下左右各 25mm
    正文: 宋体 小四号 12pt       标题: 黑体 三号/四号
    行距: 1.5 倍                首行缩进: 两字符
    页眉: 文件类型名称 黑体五号  页脚: "第 X 页" 居中
    分节: 权利要求书/说明书/摘要各自独立页眉和页码
"""

import argparse
import json
import sys
from pathlib import Path

FONT_BODY = "宋体"
FONT_BODY_ASCII = "Times New Roman"
FONT_HEADING = "黑体"
FONT_HEADING_ASCII = "Arial"


def _set_font(run, cn=FONT_BODY, ascii_=FONT_BODY_ASCII, size=None, bold=False):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    if size is None:
        size = Pt(12)
    run.font.size = size
    run.font.bold = bold
    run.font.name = ascii_
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), ascii_)
    rFonts.set(qn('w:hAnsi'), ascii_)


def _setup_page(section):
    from docx.shared import Cm
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.5))


def _set_header(section, text):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    h = section.header
    h.is_linked_to_previous = False
    p = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(text), FONT_HEADING, FONT_HEADING_ASCII, Pt(10.5))


def _set_footer(section):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    f = section.footer
    f.is_linked_to_previous = False
    p = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("第 "), size=Pt(9))
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    p.add_run()._element.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    p.add_run()._element.append(instr)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    p.add_run()._element.append(fld_end)
    _set_font(p.add_run(" 页"), size=Pt(9))


def _heading(doc, text, level=1):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    size = Pt(15) if level == 1 else Pt(14)
    _set_font(p.add_run(text), FONT_HEADING, FONT_HEADING_ASCII, size, bold=True)


def _body(doc, text):
    from docx.shared import Pt, Cm

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(0.74)
    _set_font(p.add_run(text))


def create(content_path: str, output_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Cm

    data = json.loads(Path(content_path).read_text(encoding="utf-8"))
    doc = Document()
    sec0 = doc.sections[0]
    _setup_page(sec0)
    has = False

    if data.get("claims"):
        _set_header(sec0, "权利要求书")
        _set_footer(sec0)
        _heading(doc, "权 利 要 求 书")
        for c in data["claims"]:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_after = Pt(6)
            if c.get("dependent"):
                pf.left_indent = Cm(0.74)
            r1 = p.add_run(f"{c['number']}. ")
            _set_font(r1, bold=True)
            _set_font(p.add_run(c["text"]))
        has = True

    if data.get("specification"):
        if has:
            sec = doc.add_section()
            _setup_page(sec)
        else:
            sec = sec0
        _set_header(sec, "说  明  书")
        _set_footer(sec)
        _heading(doc, "说  明  书")
        for key, title in [
            ("field", "技术领域"), ("background", "背景技术"),
            ("summary", "发明内容"), ("drawings", "附图说明"),
            ("detailed", "具体实施方式"),
        ]:
            content = data["specification"].get(key)
            if not content:
                continue
            _heading(doc, title, 2)
            for t in (content if isinstance(content, list) else [content]):
                _body(doc, t)
        has = True

    if data.get("abstract"):
        if has:
            sec = doc.add_section()
            _setup_page(sec)
        else:
            sec = sec0
        _set_header(sec, "摘    要")
        _set_footer(sec)
        _heading(doc, "摘    要")
        _body(doc, data["abstract"])
        has = True

    if not has:
        print("错误：JSON 中需要 claims / specification / abstract 至少其一", file=sys.stderr)
        sys.exit(1)

    doc.save(output_path)
    kb = Path(output_path).stat().st_size / 1024
    print(f"已生成：{output_path}（{kb:.1f} KB）")


def main():
    parser = argparse.ArgumentParser(description="从零生成标准排版的专利 .docx")
    parser.add_argument("content", help="专利内容 JSON 文件")
    parser.add_argument("output", help="输出 .docx 文件路径")
    args = parser.parse_args()

    if not Path(args.content).exists():
        print(f"错误：文件不存在：{args.content}", file=sys.stderr)
        sys.exit(1)

    create(args.content, args.output)


if __name__ == "__main__":
    main()
